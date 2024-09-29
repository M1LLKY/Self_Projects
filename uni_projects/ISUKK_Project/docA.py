import sqlite3
import docx

# Подключение к базе данных
conn = sqlite3.connect('invoices.db')
cursor = conn.cursor()

# Запрос для получения последней записи из базы данных
cursor.execute('SELECT invoice_number, date, invoice_name, employee_name, employee_position, goods_name, goods_price, goods_count, total_amount FROM invoices ORDER BY id DESC LIMIT 1')
last_invoice = cursor.fetchone()

# Извлечение данных из последней записи
invoice_number, date, invoice_name, employee_name, employee_position, goods_name, goods_price, goods_count, total_amount = last_invoice

# Открытие готового документа Word
doc = docx.Document('шаблон_товарной_накладной.docx')

# Замена переменных в документе
for paragraph in doc.paragraphs:
    if '{{INVOICE_NUMBER}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{INVOICE_NUMBER}}', invoice_number)
    if '{{DATE}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{DATE}}', date)
    if '{{INVOICE_NAME}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{INVOICE_NAME}}', invoice_name)
    if '{{EMPLOYEE_NAME}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{EMPLOYEE_NAME}}', employee_name)
    if '{{EMPLOYEE_POSITION}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{EMPLOYEE_POSITION}}', employee_position)
    if '{{COUNT_POS}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{COUNT_POS}}', str(1))  # Предполагаем, что всегда 1 позиция
    if '{{ALL_SUM}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{ALL_SUM}}', str(total_amount))
    if '{{EMPLOYEE_SHORTNAME}}' in paragraph.text:
        paragraph.text = paragraph.text.replace('{{EMPLOYEE_SHORTNAME}}', employee_name.split()[1])  # Используем фамилию как короткое имя

# Поиск существующей таблицы
table = None
for t in doc.tables:
    if t.cell(0, 0).text == 'Экскурсия':
        table = t
        break

if table is None:
    # Если таблица не найдена, создаем новую
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Экскурсия'
    hdr_cells[1].text = 'Цена'
    hdr_cells[2].text = 'Количество участников'
    hdr_cells[3].text = 'Сумма'
else:
    # Если таблица найдена, очищаем её содержимое, кроме заголовков
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)  # Удаляем все строки, кроме заголовков

# Добавление данных о товаре в таблицу
row_cells = table.add_row().cells
row_cells[0].text = goods_name
row_cells[1].text = str(goods_price)
row_cells[2].text = str(goods_count)
row_cells[3].text = str(total_amount)

# Добавление итоговой строки в таблицу
total_row_cells = table.add_row().cells
total_row_cells[0].text = 'Итого:'
total_row_cells[1].text = str(goods_price)  # Указываем цену
total_row_cells[2].text = str(goods_count)  # Указываем количество
total_row_cells[3].text = str(total_amount)  # Указываем общую сумму

# Сохранение документа
doc.save('товарная_накладная.docx')

# Закрытие соединения с базой данных
conn.close()