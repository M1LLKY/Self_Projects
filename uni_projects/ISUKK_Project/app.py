import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_from_directory, url_for
from werkzeug.exceptions import abort
import os
import pprint
from docx import Document
from datetime import datetime

app = Flask(__name__)
app.config['SECRET_KEY'] = 'some_unique_and_secret_key'

DIRECTORY = os.getcwd() + "/"
NUMBER = 0

# Временный словарь для хранения данных
TEMP_DICT = {
    "name": '',
    "passport": '',
    "pc_id": ''
}


# Функция для подключения к базе данных
def get_db_connection(db_name):
    conn = sqlite3.connect(db_name)
    conn.row_factory = sqlite3.Row
    return conn


# Функция для создания таблиц при запуске приложения
def create_tables():
    conn = get_db_connection('database.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS Contracts (
            id_contract INTEGER PRIMARY KEY,
            number TEXT,
            date TEXT,
            deal_type TEXT,
            start_price REAL,
            discount REAL,
            deal_status TEXT,
            finall_price REAL,
            id INTEGER,
            client_id INTEGER,
            employee_id INTEGER
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS Ticket (
            id INTEGER PRIMARY KEY,
            CPU TEXT,
            GPU TEXT,
            Motherboard TEXT,
            RAM TEXT,
            "Case" TEXT,
            available BOOLEAN,
            client_id INTEGER,
            price INTEGER
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS Clients (
            id_client INTEGER PRIMARY KEY,
            name TEXT,
            email TEXT,
            phone_number TEXT,
            passport TEXT
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS Employees (
            id INTEGER PRIMARY KEY,
            name TEXT,
            email TEXT,
            phone_number TEXT,
            position TEXT,
            department TEXT,
            chief_id INTEGER
        )
    ''')
    conn.commit()
    conn.close()


create_tables()


# Главная страница
@app.route('/')
def index():
    return render_template("index.html")


# Список контрактов
@app.route('/contracts')
def contracts():
    conn = get_db_connection('database.db')
    inv = conn.execute('SELECT * FROM Contracts').fetchall()
    conn.close()
    return render_template("contracts.html", contracts=inv)


# Детали одного контракта
@app.route('/contracts/<int:contract_id>')
def one_contract(contract_id):
    conn = get_db_connection('database.db')
    contract = conn.execute('SELECT * FROM Contracts WHERE id_contract=?', (contract_id,)).fetchone()
    if contract is None:
        abort(404)
    conn.close()
    return render_template("one_contract.html", contract=contract)


# Создание нового контракта
@app.route('/contracts/create_contract', methods=['GET', 'POST'])
def create_contract():
    if request.method == 'POST':
        conn = get_db_connection('database.db')
        try:
            previous_contract = conn.execute('SELECT * FROM Contracts').fetchall()[-1]
            new_id = previous_contract['id_contract'] + 1
        except IndexError:
            new_id = 0
        number = request.form.get('number')
        date = request.form.get('date')
        deal_type = request.form.get('deal_type')
        start_price = request.form.get('start_price')
        discount = request.form.get('discount')
        deal_status = request.form.get('deal_status')
        finall_price = request.form.get('finall_price')
        pc_id = request.form.get('id')
        client_id = request.form.get('client_id')
        employee_id = request.form.get('employee_id')
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO Contracts (id_contract, number, date, deal_type, start_price, discount, deal_status, finall_price, id, client_id, employee_id) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (new_id, number, date, deal_type, start_price, discount, deal_status, finall_price, pc_id,
             client_id, employee_id))
        conn.commit()
        conn.close()
        return redirect(url_for('contracts'))
    return render_template('create_contract.html')


# Список билетов
@app.route('/ticket')
def ticket():
    conn = get_db_connection('database.db')
    list_of_ticket = conn.execute('SELECT * FROM Ticket').fetchall()
    conn.close()
    return render_template("ticket.html", ticket=list_of_ticket)


# Детали одного билета
@app.route('/ticket/<int:id>')
def one_ticket(id):
    conn = get_db_connection('database.db')
    tickets = conn.execute('SELECT * FROM Ticket').fetchall()
    if id >= len(tickets):
        abort(404)
    ticket = tickets[id]
    conn.close()
    return render_template("one_ticket.html", ticket=ticket)


# Создание нового билета
@app.route('/ticket/create_ticket', methods=['GET', 'POST'])
def create_ticket():
    if request.method == 'POST':
        conn = get_db_connection('database.db')
        try:
            previous_ticket = conn.execute('SELECT * FROM Ticket').fetchall()[-1]
            new_id = previous_ticket['id'] + 1
        except IndexError:
            new_id = 0
        CPU = request.form.get('CPU')
        GPU = request.form.get('GPU')
        Motherboard = request.form.get('Motherboard')
        RAM = request.form.get('RAM')
        Case = request.form.get('Case')
        available = request.form.get('available')
        client_id = request.form.get('client_id')
        price = request.form.get('price')
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO Ticket (id, CPU, GPU, Motherboard, RAM, \"Case\", available, client_id, price) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (new_id, CPU, GPU, Motherboard, RAM, Case, available, client_id, price))
        conn.commit()
        conn.close()
        return redirect(url_for('ticket'))
    return render_template('create_ticket.html')


# Функция удаления строк
@app.route('/remove/<table>-<int:row_id>', methods=['GET', 'POST'])
def remove_row(table, row_id=-1):
    conn = get_db_connection('database.db')
    cur = conn.cursor()
    cur.execute(f'DELETE FROM {table} WHERE id=?', (row_id,))
    conn.commit()
    conn.close()
    return redirect(url_for(table))


# Список клиентов
@app.route('/clients')
def clients():
    conn = get_db_connection('database.db')
    cli = conn.execute('SELECT * FROM Clients').fetchall()
    conn.close()
    return render_template("clients.html", clients=cli)


# Детали одного клиента
@app.route('/clients/<int:id_client>')
def one_client(id_client):
    conn = get_db_connection('database.db')
    clients = conn.execute('SELECT * FROM Clients').fetchall()
    if id_client >= len(clients):
        abort(404)
    client = clients[id_client]
    conn.close()
    return render_template("one_client.html", client=client)


# Создание нового клиента
@app.route('/clients/create_client', methods=['GET', 'POST'])
def create_client():
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        phone_number = request.form.get('phone_number')
        passport = request.form.get('passport')

        # Проверка наличия ключей
        if not all([name, email, phone_number, passport]):
            flash("Не все поля формы заполнены")
            return redirect(url_for('create_client'))

        conn = get_db_connection('database.db')
        try:
            previous_client = conn.execute('SELECT * FROM Clients').fetchall()[-1]
            new_id = previous_client['id_client'] + 1
        except IndexError:
            new_id = 0

        cursor = conn.cursor()
        cursor.execute('INSERT INTO Clients (id_client, name, email, phone_number, passport) VALUES (?, ?, ?, ?, ?)',
                       (new_id, name, email, phone_number, passport))
        conn.commit()
        conn.close()
        return redirect(url_for('clients'))
    return render_template('create_client.html')


# Список сотрудников
@app.route('/employees')
def employees():
    conn = get_db_connection('database.db')
    empls = conn.execute('SELECT * FROM Employees').fetchall()
    conn.close()
    return render_template("employees.html", employees=empls)


# Детали одного сотрудника
@app.route('/employees/<int:employee_id>')
def one_employee(employee_id):
    conn = get_db_connection('database.db')
    employees = conn.execute('SELECT * FROM Employees').fetchall()
    if employee_id >= len(employees):
        abort(404)
    employee = employees[employee_id]
    conn.close()
    return render_template("one_employee.html", employee=employee)


# Создание нового сотрудника
@app.route('/employees/create_employee', methods=['GET', 'POST'])
def create_employee():
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        phone_number = request.form.get('phone_number')
        position = request.form.get('position')
        department = request.form.get('department')
        chief_id = request.form.get('chief_id')

        # Проверка наличия ключей
        if not all([name, email, phone_number, position, department, chief_id]):
            flash("Не все поля формы заполнены")
            return redirect(url_for('create_employee'))

        conn = get_db_connection('database.db')
        try:
            previous_employee = conn.execute('SELECT * FROM Employees').fetchall()[-1]
            new_id = previous_employee['id'] + 1
        except IndexError:
            new_id = 0

        cursor = conn.cursor()
        cursor.execute(
            'INSERT INTO Employees (id, name, email, phone_number, position, department, chief_id) VALUES (?, ?, ?, ?, ?, ?, ?)',
            (new_id, name, email, phone_number, position, department, chief_id))
        conn.commit()
        conn.close()
        return redirect(url_for('employees'))
    return render_template('create_employee.html')


# Функция для получения данных для отчета
def get_data_for_report():
    conn = get_db_connection('database.db')
    ticket = conn.execute('SELECT * FROM Ticket WHERE id=?', (TEMP_DICT['pc_id'],)).fetchone()
    if ticket is None:
        flash("Запись не найдена в базе данных")
        return redirect(url_for('index'))

    person = conn.execute('SELECT * FROM Employees WHERE department=?', ('Хозяйственный Отдел',)).fetchall()
    if not person:
        flash("Сотрудник не найден в базе данных")
        return redirect(url_for('index'))

    conn.close()
    id = ticket['id']
    price = ticket['price']
    client_name = TEMP_DICT['name']
    client_passport = TEMP_DICT['passport']
    return [id, price, client_name, client_passport], [person['name'], person['position']]


# Создание отчета
@app.route('/create_report', methods=['GET', 'POST'])
def create_report():
    if request.method == 'POST':
        name = request.form.get('name')
        passport = request.form.get('passport')
        pc_id = request.form.get('id')
        TEMP_DICT['name'] = name
        TEMP_DICT['passport'] = passport
        TEMP_DICT['pc_id'] = pc_id
        return redirect(url_for('save_report'))
    return render_template("create_report.html")


# Сохранение отчета
@app.route('/save_report')
def save_report():
    try:
        table_data, sub_data = get_data_for_report()
    except Exception as e:
        flash("Произошла ошибка: " + str(e))
        return redirect(url_for('index'))

    number = 0
    template_name = DIRECTORY + "docx_templates/template_report.docx"
    template_doc = Document(template_name)
    data = {}
    with open("docx_templates/report_keys.txt", 'rt', encoding='utf-8') as data_txt:
        not_so_keys = [elem.strip() for elem in data_txt.readlines()]
        for k in range(len(not_so_keys)):
            data[not_so_keys[k]] = ""
    data["{{REPORT_NUMBER}}"] = str(number)
    data["{{DATE}}"] = " ".join(list(reversed(str(datetime.today().date()).split("-"))))
    table = template_doc.tables
    table.add_row()
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = str(table_data)
    table.cell(1, 2).text = str(table_data)
    table.cell(1, 3).text = str(table_data)
    table.cell(1, 4).text = str(table_data)
    data["{{EMPLOYEE_NAME}}"] = sub_data
    data["{{EMPLOYEE_POSITION}}"] = sub_data
    for key, value in data.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, key, value)

    template_doc.save(DIRECTORY + "/reports/Договор о продаже №" + str(number) + ".docx")
    return redirect(url_for('index'))


# Список отчетов
@app.route('/reports')
def reports():
    reports_list = os.listdir(DIRECTORY + "/reports")
    return render_template("reports.html", reports=reports_list)


# Скачивание отчета
@app.route('/download_report/<string:report_name>')
def download_report(report_name):
    return send_from_directory(DIRECTORY + "/reports", report_name, as_attachment=True)


# Обработчик ошибок 404
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


# Обработчик ошибок 500
@app.errorhandler(500)
def internal_server_error(e):
    return render_template('500.html'), 500


# Функция для замены текста в документе
def replace_text(paragraph, key, value):
    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


# Функция для удаления подписей
def remove_sub_script():
    try:
        os.remove(DIRECTORY + "\\Отчет о наличии №0.docx")
    except FileNotFoundError:
        pass


if __name__ == '__main__':
    remove_sub_script()
    print(DIRECTORY)
    app.run(debug=True)